from __future__ import annotations

import csv
import io
import json
from pathlib import Path


TEXT_EXTENSIONS = {
    ".txt",
    ".csv",
    ".json",
    ".jsonl",
    ".log",
    ".sql",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".java",
    ".c",
    ".cpp",
    ".cs",
    ".go",
    ".rs",
    ".php",
    ".rb",
    ".env",
    ".ini",
    ".conf",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".css",
    ".md",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}


def supported_extensions() -> list[str]:
    return sorted(TEXT_EXTENSIONS | IMAGE_EXTENSIONS | {".pdf", ".docx"})


def extract_text(uploaded_file) -> tuple[str, str | None]:
    """
    Extract plain text from a Django uploaded file.

    Supported: PDF, DOCX, TXT/CSV/JSON/log/source/config files, and images via
    optional OCR dependencies. The caller owns security decisions after text
    extraction.
    """
    name = uploaded_file.name or "uploaded-file"
    suffix = Path(name.lower()).suffix
    raw = uploaded_file.read()
    try:
        uploaded_file.seek(0)
    except Exception:
        pass

    if suffix == ".pdf":
        return _extract_pdf(raw)
    if suffix == ".docx":
        return _extract_docx(raw)
    if suffix == ".csv":
        return _extract_csv(raw)
    if suffix in {".json", ".jsonl"}:
        return _extract_json(raw)
    if suffix in TEXT_EXTENSIONS or _looks_like_text(raw):
        return _decode_text(raw)
    if suffix in IMAGE_EXTENSIONS:
        return _extract_image_text(raw)

    supported = ", ".join(supported_extensions())
    return "", f"Unsupported file type: '{name}'. Supported extensions: {supported}."


def _extract_pdf(raw: bytes) -> tuple[str, str | None]:
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return _non_empty("\n".join(pages), "PDF appears to contain no extractable text.")
    except ImportError:
        return "", "pdfplumber is not installed."
    except Exception as exc:
        return "", f"PDF extraction failed: {exc}"


def _extract_docx(raw: bytes) -> tuple[str, str | None]:
    try:
        from docx import Document

        document = Document(io.BytesIO(raw))
        text_parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                text_parts.append(" | ".join(cell.text for cell in row.cells))
        return _non_empty("\n".join(text_parts), "DOCX file appears to be empty.")
    except ImportError:
        return "", "python-docx is not installed."
    except Exception as exc:
        return "", f"DOCX extraction failed: {exc}"


def _extract_csv(raw: bytes) -> tuple[str, str | None]:
    text, error = _decode_text(raw)
    if error:
        return text, error
    try:
        rows = list(csv.reader(io.StringIO(text)))
    except csv.Error:
        return text, None
    flattened = "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)
    return _non_empty(flattened, "CSV file is empty.")


def _extract_json(raw: bytes) -> tuple[str, str | None]:
    text, error = _decode_text(raw)
    if error:
        return text, error
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        return text, None
    return _non_empty(json.dumps(parsed, indent=2, ensure_ascii=False), "JSON file is empty.")


def _extract_image_text(raw: bytes) -> tuple[str, str | None]:
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(io.BytesIO(raw))
        return _non_empty(pytesseract.image_to_string(image), "No text could be extracted from the image.")
    except ImportError:
        return "", "pytesseract or Pillow is not installed."
    except Exception as exc:
        return "", f"Image OCR failed: {exc}"


def _decode_text(raw: bytes) -> tuple[str, str | None]:
    for encoding in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            return _non_empty(raw.decode(encoding), "Text file is empty.")
        except UnicodeDecodeError:
            continue
    return "", "Text decode failed."


def _looks_like_text(raw: bytes) -> bool:
    if not raw:
        return True
    sample = raw[:2048]
    return b"\x00" not in sample


def _non_empty(text: str, empty_message: str) -> tuple[str, str | None]:
    cleaned = (text or "").strip()
    if not cleaned:
        return "", empty_message
    return cleaned, None
